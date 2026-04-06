import base64, io, re
from docx import Document
from docx.shared import Pt, Cm, Emu, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

F   = "Times New Roman"
FS  = Pt(11)
F14 = Pt(14)
F12 = Pt(12)

# Rasmiy standartdan olingan tab stop pozitsiyasi: 4228 twips = 7.46 cm
TAB_POS_TWIPS = 4228
# Mehnat faoliyati: hanging indent = 942340 EMU ≈ 2.62 cm
WORK_INDENT = Cm(2.62)
# Label qatorlar orasidagi masofa: 101600 EMU = 8pt
LABEL_SPACE = Pt(8)
# Mehnat faoliyati qatorlari orasidagi masofa: 50800 EMU = 4pt
WORK_SPACE = Pt(4)


# ──────────────────────────────────────────────────────────────
#  LOTIN → KIRILL TRANSLITERATSIYA
# ──────────────────────────────────────────────────────────────
def _is_cyrillic_char(c: str) -> bool:
    return '\u0400' <= c <= '\u04FF'

def _has_cyrillic(text: str) -> bool:
    return any(_is_cyrillic_char(c) for c in text)

def _normalize_apostrophe(text: str) -> str:
    for ch in ['\u2018', '\u2019', '\u02BC', '\u02BB', '\u0060', '\u00B4',
               '\u044A', '\u044B', '\u042A', '\u0027']:
        text = text.replace(ch, "'")
    return text

def _word_lat_to_cyr(word: str) -> str:
    if _has_cyrillic(word):
        return word

    DIGRAPHS = [
        ('SH', 'Ш'), ('CH', 'Ч'), ('NG', 'НГ'), ('TS', 'Ц'),
        ('YA', 'Я'), ('YO', 'Ё'), ('YU', 'Ю'), ('YE', 'Е'),
        ('Sh', 'Ш'), ('Ch', 'Ч'), ('Ng', 'НГ'), ('Ts', 'Ц'),
        ('Ya', 'Я'), ('Yo', 'Ё'), ('Yu', 'Ю'), ('Ye', 'Е'),
        ('sh', 'ш'), ('ch', 'ч'), ('ng', 'нг'), ('ts', 'ц'),
        ('ya', 'я'), ('yo', 'ё'), ('yu', 'ю'), ('ye', 'е'),
    ]
    APOSTROPHE_PAIRS = [
        ("O'", 'Ў'), ("G'", 'Ғ'),
        ("o'", 'ў'), ("g'", 'ғ'),
    ]
    SINGLES = {
        'A': 'А', 'B': 'Б', 'D': 'Д', 'E': 'Е', 'F': 'Ф', 'G': 'Г',
        'H': 'Ҳ', 'I': 'И', 'J': 'Ж', 'K': 'К', 'L': 'Л', 'M': 'М',
        'N': 'Н', 'O': 'О', 'P': 'П', 'Q': 'Қ', 'R': 'Р', 'S': 'С',
        'T': 'Т', 'U': 'У', 'V': 'В', 'X': 'Х', 'Y': 'Й', 'Z': 'З',
        'a': 'а', 'b': 'б', 'd': 'д', 'e': 'е', 'f': 'ф', 'g': 'г',
        'h': 'ҳ', 'i': 'и', 'j': 'ж', 'k': 'к', 'l': 'л', 'm': 'м',
        'n': 'н', 'o': 'о', 'p': 'п', 'q': 'қ', 'r': 'р', 's': 'с',
        't': 'т', 'u': 'у', 'v': 'в', 'x': 'х', 'y': 'й', 'z': 'з',
        "'": 'ъ',
    }

    s = _normalize_apostrophe(word)
    for lat, cyr in APOSTROPHE_PAIRS:
        s = s.replace(lat, cyr)
    for lat, cyr in DIGRAPHS:
        s = s.replace(lat, cyr)

    out = []
    for i, ch in enumerate(s):
        if ch in SINGLES:
            if ch == 'E':
                prev = s[i-1] if i > 0 else ''
                if i == 0 or (prev not in 'aeiouAEIOUаеёиоуыэюяўАЕЁИОУЫЭЮЯЎ'
                              and not _is_cyrillic_char(prev)
                              and prev not in 'аеёиоуыэюяўbdfghjklmnpqrstvxyz'):
                    out.append('Э')
                else:
                    out.append(SINGLES[ch])
            elif ch == 'e':
                prev = s[i-1] if i > 0 else ''
                if i == 0 or prev in ' \t\n\r-–—.,;:!?()"«»':
                    out.append('э')
                else:
                    out.append(SINGLES[ch])
            else:
                out.append(SINGLES[ch])
        else:
            out.append(ch)
    return ''.join(out)


def lat_to_cyr(text: str) -> str:
    if not text:
        return text
    cyr_count = sum(1 for c in text if _is_cyrillic_char(c))
    lat_count = sum(1 for c in text if c.isalpha() and not _is_cyrillic_char(c))
    if cyr_count > lat_count:
        return text
    parts = re.split(r'(\s+)', text)
    return ''.join(_word_lat_to_cyr(p) if p.strip() else p for p in parts)


def apply_script(data: dict, script: str) -> dict:
    if script != 'cyr':
        return data
    TEXT_FIELDS = [
        'fullname', 'birthplace', 'nationality', 'party',
        'edu_level', 'university', 'speciality',
        'science_degree', 'science_title', 'military_rank',
        'langs', 'awards', 'departmental_awards', 'deputy',
        'address', 'current_job', 'job_year',
    ]
    result = dict(data)
    for field in TEXT_FIELDS:
        val = result.get(field)
        if isinstance(val, str):
            result[field] = lat_to_cyr(val)
        elif isinstance(val, list):
            result[field] = [lat_to_cyr(v) if isinstance(v, str) else v for v in val]
    wh = result.get('work_history', [])
    result['work_history'] = [
        {k: lat_to_cyr(v) if isinstance(v, str) else v for k, v in w.items()}
        for w in wh if isinstance(w, dict)
    ]
    rels = result.get('relatives', [])
    result['relatives'] = [
        {k: lat_to_cyr(v) if isinstance(v, str) else v for k, v in r.items()}
        for r in rels if isinstance(r, dict)
    ]
    return result


# ──────────────────────────────────────────────────────────────
#  FONT / PARAGRAF YORDAMCHILARI
# ──────────────────────────────────────────────────────────────
def _set_rfonts(r_pr, name=F):
    """Run ning barcha font xossalarini to'g'ri o'rnatish"""
    rFonts = r_pr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        r_pr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)
    rFonts.set(qn('w:eastAsia'), name)


def _font(run, bold=False, size=None):
    run.font.name = F
    _set_rfonts(run._element.get_or_add_rPr())
    run.font.size = size or FS
    run.font.bold = bold


def _run(p, text, bold=False, size=None):
    r = p.add_run(text)
    _font(r, bold=bold, size=size)
    return r


def _add_tab_stop(p, pos_twips):
    """Paragrafga tab stop qo'shish (twips da)"""
    pPr = p._element.get_or_add_pPr()
    tabs_el = pPr.find(qn('w:tabs'))
    if tabs_el is None:
        tabs_el = OxmlElement('w:tabs')
        pPr.append(tabs_el)
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'left')
    tab.set(qn('w:pos'), str(pos_twips))
    tabs_el.append(tab)


# ──────────────────────────────────────────────────────────────
#  FLOATING RASM (ANCHOR) — rasmiy standartga mos
# ──────────────────────────────────────────────────────────────
def add_floating_image(paragraph, image_bytes, width_cm=3.0, height_cm=4.0):
    """
    Rasm yuqori-o'ng burchakda anchor tarzda joylashtiriladi.
    Rasmiy standart: 3×4 sm, wrapSquare, margin-right=margin-top ga yaqin.
    """
    run = paragraph.add_run()
    img_stream = io.BytesIO(image_bytes)
    inline_shape = run.add_picture(img_stream, width=Cm(width_cm), height=Cm(height_cm))

    inline = run._element.xpath('.//wp:inline')
    if not inline:
        return inline_shape
    inline = inline[0]

    cx = str(int(Cm(width_cm)))
    cy = str(int(Cm(height_cm)))

    blip_list = inline.xpath('.//a:blip')
    r_embed = ""
    if blip_list:
        r_embed = blip_list[0].get(qn('r:embed')) or ""

    # Rasmiy namunadan: rasm sahifaning yuqori-o'ng burchagida,
    # margin ga nisbatan right-aligned, top-aligned
    anchor_xml = f'''<wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
    xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"
    xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    behindDoc="0" distT="0" distB="0" distL="114300" distR="0"
    simplePos="0" locked="0" layoutInCell="1" allowOverlap="1" relativeHeight="2">
  <wp:simplePos x="0" y="0"/>
  <wp:positionH relativeFrom="margin">
    <wp:align>right</wp:align>
  </wp:positionH>
  <wp:positionV relativeFrom="margin">
    <wp:align>top</wp:align>
  </wp:positionV>
  <wp:extent cx="{cx}" cy="{cy}"/>
  <wp:effectExtent l="0" t="0" r="0" b="0"/>
  <wp:wrapSquare wrapText="bothSides"/>
  <wp:docPr id="1" name="Photo1"/>
  <wp:cNvGraphicFramePr>
    <a:graphicFrameLocks noChangeAspect="1"/>
  </wp:cNvGraphicFramePr>
  <a:graphic>
    <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
      <pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">
        <pic:nvPicPr>
          <pic:cNvPr id="1" name="Photo1"/>
          <pic:cNvPicPr>
            <a:picLocks noChangeAspect="1" noChangeArrowheads="1"/>
          </pic:cNvPicPr>
        </pic:nvPicPr>
        <pic:blipFill>
          <a:blip r:embed="{r_embed}"/>
          <a:stretch><a:fillRect/></a:stretch>
        </pic:blipFill>
        <pic:spPr bwMode="auto">
          <a:xfrm>
            <a:off x="0" y="0"/>
            <a:ext cx="{cx}" cy="{cy}"/>
          </a:xfrm>
          <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
        </pic:spPr>
      </pic:pic>
    </a:graphicData>
  </a:graphic>
</wp:anchor>'''

    new_anchor = etree.fromstring(anchor_xml)
    inline.getparent().replace(inline, new_anchor)
    return inline_shape


# ──────────────────────────────────────────────────────────────
#  JADVAL YORDAMCHILARI (2-sahifa: qarindoshlar)
# ──────────────────────────────────────────────────────────────
def _cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    cb = OxmlElement("w:tcBorders")
    for s in ["top", "left", "bottom", "right"]:
        b = OxmlElement(f"w:{s}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "000000")
        cb.append(b)
    tcPr.append(cb)


def _cell_para(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=None):
    p = cell.paragraphs[0]
    p.clear()
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.line_spacing = Pt(13)
    _run(p, text or "—", bold=bold, size=size or FS)


# ──────────────────────────────────────────────────────────────
#  ASOSIY GENERATOR
# ──────────────────────────────────────────────────────────────
def generate(data: dict, output_path: str, script: str = 'lat'):
    data = apply_script(data, script)
    doc = Document()

    # ─── Rasmiy sahifa sozlamalari ───
    # Izoh: yuqoridan 1.5sm, pastdan 1sm, o'ngdan 1sm, chapdan 2sm
    # (rasmiy namunada left=2.69sm bo'lgan — lekin izoh 2sm deydi,
    #  biz rasmiy amaliy parametrni ishlatamiz)
    sec = doc.sections[0]
    sec.page_width    = Cm(21)
    sec.page_height   = Cm(29.7)
    sec.top_margin    = Cm(1.5)
    sec.bottom_margin = Cm(1.0)
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(1.0)

    # Normal uslub
    doc.styles["Normal"].font.name = F
    doc.styles["Normal"].font.size = FS
    _set_rfonts(doc.styles["Normal"]._element.get_or_add_rPr())

    fullname    = data.get("fullname", "")
    job_year    = data.get("job_year", "")
    current_job = data.get("current_job", "")
    photo_b64   = data.get("photo_base64", "")

    bd = data.get("birthdate", "")
    if bd and "-" in bd:
        y, m, d = bd.split("-")
        bd = f"{d}.{m}.{y}"
    langs = data.get("langs", [])
    langs_str = ", ".join(langs) if isinstance(langs, list) else str(langs)

    IS_CYR = script == "cyr"
    YOQ = "йўқ" if IS_CYR else "yo'q"
    def L(lat, cyr): return cyr if IS_CYR else lat

    # ═══════════════════════════════════════════════════════════
    # 1-SAHIFA
    # ═══════════════════════════════════════════════════════════

    # ─── SARLAVHA: МАЪЛУМОТНОМА — 14pt bold, markazda ───
    p0 = doc.add_paragraph()
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before = Pt(0)
    p0.paragraph_format.space_after  = Pt(0)
    # Birinchi paragrafga floating rasmni biriktirish
    if photo_b64:
        try:
            img_bytes = base64.b64decode(photo_b64.split(",")[-1])
            add_floating_image(p0, img_bytes, width_cm=3.0, height_cm=4.0)
        except Exception as e:
            pass
    _run(p0, L("MA'LUMOTNOMA", "МАЪЛУМОТНОМА"), bold=True, size=F14)

    # ─── F.I.Sh. — 14pt bold, markazda ───
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after  = Pt(0)
    _run(p1, fullname, bold=True, size=F14)

    # ─── Bo'sh qator (namunada bor) ───
    doc.add_paragraph().paragraph_format.space_before = Pt(0)

    # ─── Joriy lavozim: sana va ish joyi ───
    if job_year:
        p_jy = doc.add_paragraph()
        p_jy.paragraph_format.space_before = Pt(0)
        p_jy.paragraph_format.space_after  = Pt(0)
        _run(p_jy, job_year, size=FS)
    if current_job:
        p_cj = doc.add_paragraph()
        p_cj.paragraph_format.space_before = Pt(0)
        p_cj.paragraph_format.space_after  = Pt(0)
        _run(p_cj, current_job, size=FS)

    # ═══════════════════════════════════════════════════════════
    # 2-USTUNLI MA'LUMOT BLOKLARI (tab stop = 4228 twips = 7.46 cm)
    # Rasmiy namuna: label = bold 11pt, space_before=8pt
    #                value = regular 11pt, space_before=0
    # ═══════════════════════════════════════════════════════════

    def add_label_row(l1, l2=None):
        """Tepadagi qator — bold labellar (2 ustunli)"""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = LABEL_SPACE
        p.paragraph_format.space_after  = Pt(0)
        _add_tab_stop(p, TAB_POS_TWIPS)
        _run(p, l1, bold=True)
        if l2:
            _run(p, "\t")
            _run(p, l2, bold=True)
        return p

    def add_value_row(v1, v2=None):
        """Pastdagi qator — oddiy qiymatlar"""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        _add_tab_stop(p, TAB_POS_TWIPS)
        _run(p, v1 or "—")
        if v2 is not None:
            _run(p, "\t")
            _run(p, v2 or "—")
        return p

    def add_inline_pair(label, val):
        """Bir qatorda label + tab + val"""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = LABEL_SPACE
        p.paragraph_format.space_after  = Pt(0)
        _add_tab_stop(p, TAB_POS_TWIPS)
        _run(p, label, bold=True)
        _run(p, "\t")
        _run(p, val or "—")
        return p

    def add_full_label(label):
        """To'liq kenglikdagi bold label"""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = LABEL_SPACE
        p.paragraph_format.space_after  = Pt(0)
        _run(p, label, bold=True)
        return p

    def add_full_value(val, hanging=None):
        """To'liq kenglikdagi qiymat (ixtiyoriy hanging indent bilan)"""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        if hanging:
            p.paragraph_format.left_indent      = hanging
            p.paragraph_format.first_line_indent = -hanging
        _run(p, val or "—")
        return p

    # ─── Qatorlarni ketma-ket chiqaramiz ───

    # Tug'ilgan yili / Tug'ilgan joyi
    add_label_row(L("Tug'ilgan yili:", "Туғилган йили:"),
                  L("Tug'ilgan joyi:", "Туғилган жойи:"))
    add_value_row(bd, data.get("birthplace", ""))

    # Millati / Partiyaviyligi
    add_label_row(L("Millati:", "Миллати:"),
                  L("Partiyaviyligi:", "Партиявийлиги:"))
    add_value_row(data.get("nationality", "o'zbek"), data.get("party", YOQ))

    # Ma'lumoti / Tamomlagan
    add_label_row(L("Ma'lumoti:", "Маълумоти:"),
                  L("Tamomlagan:", "Тамомлаган:"))
    add_value_row(data.get("edu_level", ""), data.get("university", ""))

    # Mutaxassisligi (bir qatorda)
    add_inline_pair(L("Ma'lumoti bo'yicha mutaxassisligi:",
                      "Маълумоти бўйича мутахассислиги:"),
                    data.get("speciality", "") or "—")

    # Ilmiy darajasi / Ilmiy unvoni
    add_label_row(L("Ilmiy darajasi:", "Илмий даражаси:"),
                  L("Ilmiy unvoni:", "Илмий унвони:"))
    add_value_row(data.get("science_degree", YOQ),
                  data.get("science_title", YOQ))

    # Chet tillari / Harbiy unvoni
    add_label_row(L("Qaysi chet tillarini biladi:", "Қайси чет тилларини билади:"),
                  L("Harbiy (maxsus) unvoni:", "Ҳарбий (махсус) унвони:"))
    add_value_row(langs_str or YOQ, data.get("military_rank", YOQ))

    # Davlat mukofotlari (to'liq kenglikda)
    add_full_label(L("Davlat mukofotlari va premiyalari bilan taqdirlangan (qanaqa):",
                     "Давлат мукофотлари ва премиялари билан тақдирланганми (қанақа):"))
    add_full_value(data.get("awards", YOQ))

    # Idoraviy mukofotlar (to'liq kenglikda)
    add_full_label(L("Idoraviy mukofotlar bilan taqdirlangan (qanaqa):",
                     "Идоравий мукофотлар билан тақдирланганми (қанақа):"))
    add_full_value(data.get("departmental_awards", YOQ))

    # Deputatlik (to'liq kenglikda)
    add_full_label(L("Xalq deputatlari, respublika, viloyat, shahar va tuman Kengashi deputatimi yoki boshqa",
                     "Халқ депутатлари, республика, вилоят, шаҳар ва туман Кенгаши депутатими ёки бошқа"))
    # 2-qator — davomi
    p_dep2 = doc.add_paragraph()
    p_dep2.paragraph_format.space_before = Pt(0)
    p_dep2.paragraph_format.space_after  = Pt(0)
    _run(p_dep2, L("saylanadigan organlarning a'zosimi (to'liq ko'rsatilishi lozim):",
                    "сайланадиган органларнинг аъзосими (тўлиқ кўрсатилиши лозим):"), bold=True)

    dep_val = data.get("deputy", YOQ)
    add_full_value(dep_val, hanging=Cm(2.49))

    # ═══════════════════════════════════════════════════════════
    # MEHNAT FAOLIYATI — 14pt bold, markazda
    # Qatorlar: hanging indent = 2.62 cm, space_before = 4pt
    # ═══════════════════════════════════════════════════════════
    p_mf = doc.add_paragraph()
    p_mf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_mf.paragraph_format.space_before = Pt(6)
    p_mf.paragraph_format.space_after  = Pt(0)
    _run(p_mf, L("MEHNAT FAOLIYATI", "МЕҲНАТ ФАОЛИЯТИ"), bold=True, size=F14)

    for w in data.get("work_history", []):
        if isinstance(w, dict):
            f, t = w.get("from", ""), w.get("to", "")
            org, pos = w.get("org", ""), w.get("pos", "")
            hv_variants = ("h.v.", "х.в.", "х.в", "ҳ.в.", "ҳ.в", "hozirgi vaqtgacha")
            if f and t and t.lower() not in hv_variants:
                prefix = L(f"{f}-{t} yy. - ", f"{f}-{t} йй. - ")
            elif f and t:
                prefix = L(f"{f} y. -  h.v.  - ", f"{f} й. -  ҳ.в.  - ")
            elif f:
                prefix = L(f"{f} y. - ", f"{f} й. - ")
            else:
                prefix = ""
            body = org
            if pos:
                body += f" {pos}"
        else:
            prefix, body = "", str(w)

        full_line = prefix + body
        if full_line.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_before = WORK_SPACE
            p.paragraph_format.space_after  = Pt(0)
            p.paragraph_format.left_indent      = WORK_INDENT
            p.paragraph_format.first_line_indent = -WORK_INDENT
            _run(p, full_line)

    # ═══ Telefon (ixtiyoriy) ═══
    phones = data.get("phones", {})
    tel = []
    if phones.get("me"):     tel.append(phones["me"])
    if phones.get("father"): tel.append(L("ota: ", "ота: ") + phones["father"])
    if phones.get("mother"): tel.append(L("ona: ", "она: ") + phones["mother"])
    if tel:
        p_tel = doc.add_paragraph()
        p_tel.paragraph_format.space_before = Pt(10)
        p_tel.paragraph_format.space_after  = Pt(0)
        _run(p_tel, L("Tel.: ", "Тел.: ") + "     ".join(tel), bold=True)

    # ═══════════════════════════════════════════════════════════
    # 2-SAHIFA: QARINDOSHLAR JADVALI — 12pt bold sarlavha
    # ═══════════════════════════════════════════════════════════
    doc.add_page_break()

    p_rel_title = doc.add_paragraph()
    p_rel_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_rel_title.paragraph_format.space_before = Pt(0)
    p_rel_title.paragraph_format.space_after  = Pt(0)
    _run(p_rel_title,
         fullname + L("ning yaqin qarindoshlari haqida",
                      "нинг яқин қариндошлари ҳақида"),
         bold=True, size=F12)

    p_rel_sub = doc.add_paragraph()
    p_rel_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_rel_sub.paragraph_format.space_before = Pt(0)
    p_rel_sub.paragraph_format.space_after  = Pt(6)
    _run(p_rel_sub, L("MA'LUMOT", "МАЪЛУМОТ"), bold=True, size=F12)

    # ─── Jadval: 5 ustunli ───
    relatives = data.get("relatives", [])
    if relatives:
        rt = doc.add_table(rows=1, cols=5)
        rt.style = "Table Grid"
        rt.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Rasmiy namunadan: col widths
        widths = [Cm(2.22), Cm(3.81), Cm(3.49), Cm(5.04), Cm(4.00)]
        for i, w in enumerate(widths):
            rt.columns[i].width = w

        headers_lat = ["Qarindosh-\nligi", "Familiyasi, ismi\nva otasining ismi",
                       "Tug'ilgan yili\nva joyi", "Ish joyi va lavozimi", "Turar joyi"]
        headers_cyr = ["Қариндош-\nлиги", "Фамилияси, исми\nва отасининг исми",
                       "Туғилган йили\nва жойи", "Иш жойи ва лавозими", "Турар жойи"]
        headers = headers_cyr if IS_CYR else headers_lat

        for i, h in enumerate(headers):
            _cell_borders(rt.cell(0, i))
            _cell_para(rt.cell(0, i), h, bold=True,
                       align=WD_ALIGN_PARAGRAPH.CENTER)

        for rel in relatives:
            row = rt.add_row()
            byear = rel.get("byear", "")
            bplace = rel.get("bplace", "")
            if byear and bplace:
                birth = L(f"{byear} yil,\n{bplace}", f"{byear} йил,\n{bplace}")
            else:
                birth = byear or bplace
            vals = [rel.get("rel", ""), rel.get("fio", ""),
                    birth, rel.get("job", ""), rel.get("addr", "")]
            for i, v in enumerate(vals):
                _cell_borders(row.cells[i])
                if i == 0:
                    _cell_para(row.cells[i], v, bold=True,
                               align=WD_ALIGN_PARAGRAPH.CENTER)
                elif i == 2:
                    _cell_para(row.cells[i], v,
                               align=WD_ALIGN_PARAGRAPH.CENTER)
                else:
                    _cell_para(row.cells[i], v)

    doc.save(output_path)
